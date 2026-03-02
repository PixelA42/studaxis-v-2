"""
Studaxis — Export Helpers
Generate quiz/lesson content as .docx (Word) or .pdf files.
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


# ── Word (.docx) ─────────────────────────────────────────────────────────────

def quiz_to_docx(quiz_data: dict) -> bytes:
    """Return a .docx file as bytes for a generated quiz."""
    doc = Document()

    # Title
    title = doc.add_heading(quiz_data.get("quiz_title", "Quiz"), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x20, 0x3c)

    # Subtitle row
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Topic: {quiz_data.get('topic', '')}   |   "
                f"Difficulty: {quiz_data.get('difficulty', '').capitalize()}")
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph()  # spacer

    questions = quiz_data.get("questions", [])
    for i, q in enumerate(questions, 1):
        # Question text
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{i}. {q.get('question', '')}")
        q_run.bold = True
        q_run.font.size = Pt(12)

        # Options
        for opt in q.get("options", []):
            opt_para = doc.add_paragraph(style="List Bullet")
            opt_para.add_run(opt).font.size = Pt(11)

        # Answer + explanation
        answer = q.get("answer", "")
        explanation = q.get("explanation", "")
        ans_para = doc.add_paragraph()
        ans_run = ans_para.add_run(f"✔ Answer: {answer}")
        ans_run.bold = True
        ans_run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
        ans_run.font.size = Pt(11)

        if explanation:
            exp_para = doc.add_paragraph()
            exp_run = exp_para.add_run(f"💡 {explanation}")
            exp_run.italic = True
            exp_run.font.color.rgb = RGBColor(0x52, 0x79, 0x6f)
            exp_run.font.size = Pt(10)

        doc.add_paragraph()  # spacer between questions

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def notes_to_docx(notes_data: dict) -> bytes:
    """Return a .docx file as bytes for lesson study notes."""
    doc = Document()

    title = doc.add_heading(notes_data.get("title", "Study Notes"), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x20, 0x3c)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Grade: {notes_data.get('grade_level', '')}")
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    doc.add_paragraph()

    # Key concepts
    doc.add_heading("Key Concepts", level=2)
    for concept in notes_data.get("key_concepts", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(concept).font.size = Pt(11)
    doc.add_paragraph()

    # Summary
    doc.add_heading("Summary", level=2)
    summary_para = doc.add_paragraph(notes_data.get("summary", ""))
    summary_para.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # Fun fact
    fun_fact = notes_data.get("fun_fact", "")
    if fun_fact:
        doc.add_heading("Fun Fact", level=2)
        ff = doc.add_paragraph()
        ff_run = ff.add_run(f"🌟 {fun_fact}")
        ff_run.italic = True
        ff_run.font.color.rgb = RGBColor(0x52, 0x79, 0x6f)
        ff_run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF ───────────────────────────────────────────────────────────────────────
# fpdf2 2.x handles UTF-8 natively — no encoding gymnastics needed.
# Always reset X to l_margin before multi_cell to guarantee full usable width.

def _clean(text: str) -> str:
    """Replace common Unicode symbols that core PDF fonts can't render."""
    replacements = {
        "\u2705": "[OK]", "\u274c": "[X]", "\u2714": "[OK]", "\u2718": "[X]",
        "\u2019": "'",    "\u2018": "'",  "\u201c": '"',   "\u201d": '"',
        "\u2013": "-",    "\u2014": "--", "\u2026": "...", "\u00b7": "*",
        "\u2022": "*",    "\u25cf": "*",  "\u2665": "<3",  "\u2764": "<3",
        "\u00a0": " ",
    }
    for ch, rep in replacements.items():
        text = text.replace(ch, rep)
    # Drop any remaining non-latin-1 characters
    return text.encode("latin-1", "replace").decode("latin-1")


class _PDF(FPDF):
    def header(self):
        pass

    def mc(self, text: str, h: float = 6, **kwargs):
        """multi_cell that always starts at the left margin."""
        self.set_x(self.l_margin)
        self.multi_cell(0, h, _clean(text), **kwargs)

    def mc_indent(self, text: str, indent: float = 8, h: float = 6, **kwargs):
        """multi_cell with a left indent, with safe remaining width."""
        self.set_x(self.l_margin + indent)
        usable_w = self.w - self.r_margin - (self.l_margin + indent)
        self.multi_cell(usable_w, h, _clean(text), **kwargs)


def quiz_to_pdf(quiz_data: dict) -> bytes:
    """Return a PDF file as bytes for a generated quiz."""
    pdf = _PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(26, 32, 60)
    pdf.mc(quiz_data.get("quiz_title", "Quiz"), h=10, align="C")

    # Subtitle
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 116, 139)
    subtitle = (f"Topic: {quiz_data.get('topic', '')}   |   "
                f"Difficulty: {quiz_data.get('difficulty', '').capitalize()}")
    pdf.mc(subtitle, h=6, align="C")
    pdf.ln(6)

    questions = quiz_data.get("questions", [])
    for i, q in enumerate(questions, 1):
        # Question
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(26, 32, 60)
        pdf.mc(f"Q{i}. {q.get('question', '')}", h=7)
        pdf.ln(1)

        # Options — indented
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(45, 51, 74)
        for opt in q.get("options", []):
            pdf.mc_indent(opt, indent=8, h=6)

        # Answer
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(22, 163, 74)
        pdf.mc(f"Answer: {q.get('answer', '')}", h=6)

        # Explanation
        explanation = q.get("explanation", "")
        if explanation:
            pdf.set_font("Helvetica", "I", 10)
            pdf.set_text_color(82, 121, 111)
            pdf.mc(f"Explanation: {explanation}", h=6)

        pdf.ln(4)

    return bytes(pdf.output())


def notes_to_pdf(notes_data: dict) -> bytes:
    """Return a PDF file as bytes for lesson study notes."""
    pdf = _PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(26, 32, 60)
    pdf.mc(notes_data.get("title", "Study Notes"), h=10, align="C")

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 116, 139)
    pdf.mc(f"Grade: {notes_data.get('grade_level', '')}", h=6, align="C")
    pdf.ln(6)

    # Key concepts
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(26, 32, 60)
    pdf.mc("Key Concepts", h=8)
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(45, 51, 74)
    for concept in notes_data.get("key_concepts", []):
        pdf.mc_indent(f"* {concept}", indent=6, h=6)
    pdf.ln(4)

    # Summary
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(26, 32, 60)
    pdf.mc("Summary", h=8)
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(45, 51, 74)
    pdf.mc(notes_data.get("summary", ""), h=6)
    pdf.ln(4)

    # Fun fact
    fun_fact = notes_data.get("fun_fact", "")
    if fun_fact:
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(26, 32, 60)
        pdf.mc("Fun Fact", h=8)
        pdf.set_font("Helvetica", "I", 11)
        pdf.set_text_color(82, 121, 111)
        pdf.mc(f"* {fun_fact}", h=6)

    return bytes(pdf.output())
